# ============================================================
# converter.py
# PDF to Word Converter using
# pdfplumber + PyMuPDF + python-docx
# ============================================================

import os
import io
import fitz                     # PyMuPDF
import pdfplumber

from pdf2docx import Converter
from PIL import Image

from docx import Document
from docx.shared import Inches
from docx.shared import Pt

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

from docx.oxml.ns import qn


class PDFConverter:

    def __init__(self):

        self.document = None
        self.progress_callback = None

    # ----------------------------------------------------
    # Progress Callback
    # ----------------------------------------------------

    def set_progress_callback(self, callback):
        self.progress_callback = callback

    def update_progress(self, value):

        if self.progress_callback:
            self.progress_callback(value)

    # ----------------------------------------------------
    # Create Output Folder
    # ----------------------------------------------------

    def create_output_folder(self, output_file):

        folder = os.path.dirname(output_file)

        if folder == "":
            folder = os.getcwd()

        if not os.path.exists(folder):
            os.makedirs(folder)

        image_folder = os.path.join(folder, "Images")

        if not os.path.exists(image_folder):
            os.makedirs(image_folder)

        return image_folder

    # ----------------------------------------------------
    # Initialize Word Document
    # ----------------------------------------------------

    def create_document(self):

        self.document = Document()

        section = self.document.sections[0]

        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

        style = self.document.styles["Normal"]

        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(11)

    # ----------------------------------------------------
    # Add Heading
    # ----------------------------------------------------

    def add_heading(self, text):

        heading = self.document.add_heading(level=1)

        run = heading.add_run(text)

        run.bold = True

        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ----------------------------------------------------
    # Add Paragraph
    # ----------------------------------------------------

    def add_paragraph(self, text):

        if text.strip() == "":
            return

        para = self.document.add_paragraph()

        para.style = self.document.styles["Normal"]

        para.add_run(text)

    # ----------------------------------------------------
    # Save Extracted Image
    # ----------------------------------------------------

    def save_image(self, image_bytes, image_path):

        image = Image.open(io.BytesIO(image_bytes))

        image.save(image_path)

    # ----------------------------------------------------
    # Insert Image into Word
    # ----------------------------------------------------

    def add_image_to_document(self, image_path):

        try:

            self.document.add_picture(
                image_path,
                width=Inches(5.5)
            )

        except Exception:

            pass
        
        
    # ============================================================
    # Extract Text using pdfplumber
    # ============================================================

    def extract_text(self, pdf_path):

        text_data = []

        with pdfplumber.open(pdf_path) as pdf:

            total_pages = len(pdf.pages)

            for page_number, page in enumerate(pdf.pages):

                try:

                    page_text = page.extract_text(
                        x_tolerance=2,
                        y_tolerance=3
                    )

                    if page_text:

                        lines = page_text.split("\n")

                        paragraph = ""

                        for line in lines:

                            line = line.strip()

                            if line == "":

                                if paragraph:

                                    text_data.append(paragraph)
                                    paragraph = ""

                            else:

                                paragraph += line + " "

                        if paragraph:
                            text_data.append(paragraph)

                    progress = int(((page_number + 1) / total_pages) * 30)

                    self.update_progress(progress)

                except Exception as e:

                    print(f"Text Extraction Error : {e}")

        return text_data


    # ============================================================
    # Write Text to Word
    # ============================================================

    def write_text(self, text_data):

        self.add_heading("PDF Content")

        for paragraph in text_data:

            self.add_paragraph(paragraph)


    # ============================================================
    # Extract Tables
    # ============================================================

    def extract_tables(self, pdf_path):

        tables = []

        with pdfplumber.open(pdf_path) as pdf:

            total_pages = len(pdf.pages)

            for page_number, page in enumerate(pdf.pages):

                try:

                    page_tables = page.extract_tables()

                    if page_tables:

                        for table in page_tables:

                            tables.append(table)

                    progress = 30 + int(
                        ((page_number + 1) / total_pages) * 20
                    )

                    self.update_progress(progress)

                except Exception as e:

                    print(f"Table Extraction Error : {e}")

        return tables


    # ============================================================
    # Write Tables to Word
    # ============================================================

    def write_tables(self, tables):

        if not tables:
            return

        self.document.add_page_break()

        self.add_heading("Tables")

        for table_data in tables:

            if not table_data:
                continue

            rows = len(table_data)

            cols = max(len(row) for row in table_data if row)

            table = self.document.add_table(
                rows=rows,
                cols=cols
            )

            table.style = "Table Grid"

            for i, row in enumerate(table_data):

                if row is None:
                    continue

                for j, cell in enumerate(row):

                    if cell is None:
                        cell = ""

                    table.cell(i, j).text = str(cell)

            self.document.add_paragraph()
            
    # ============================================================
    # Extract Images using PyMuPDF
    # ============================================================

    def extract_images(self, pdf_path, image_folder):

        image_files = []

        try:

            pdf = fitz.open(pdf_path)

            total_pages = len(pdf)

            image_count = 1

            for page_index in range(total_pages):

                page = pdf.load_page(page_index)

                images = page.get_images(full=True)

                for img in images:

                    xref = img[0]

                    try:

                        base_image = pdf.extract_image(xref)

                        image_bytes = base_image["image"]

                        ext = base_image["ext"]

                        image_name = f"image_{image_count}.{ext}"

                        image_path = os.path.join(
                            image_folder,
                            image_name
                        )

                        with open(image_path, "wb") as f:
                            f.write(image_bytes)

                        image_files.append(image_path)

                        image_count += 1

                    except Exception as e:

                        print(f"Image Error : {e}")

                progress = 50 + int(
                    ((page_index + 1) / total_pages) * 20
                )

                self.update_progress(progress)

            pdf.close()

        except Exception as e:

            print(f"Image Extraction Failed : {e}")

        return image_files


    # ============================================================
    # Insert Images into Word
    # ============================================================

    def write_images(self, image_files):

        if not image_files:
            return

        self.document.add_page_break()

        self.add_heading("Extracted Images")

        for image in image_files:

            try:

                self.document.add_picture(
                    image,
                    width=Inches(5.8)
                )

                caption = self.document.add_paragraph()

                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

                caption.add_run(
                    os.path.basename(image)
                ).italic = True

                self.document.add_paragraph()

            except Exception as e:

                print(f"Word Image Error : {e}")


    # ============================================================
    # Image Statistics
    # ============================================================

    def get_image_count(self, pdf_path):

        count = 0

        try:

            pdf = fitz.open(pdf_path)

            for page in pdf:

                count += len(page.get_images(full=True))

            pdf.close()

        except Exception:

            pass

        return count


    # ============================================================
    # Clean Images Folder
    # ============================================================

    def clear_image_folder(self, folder):

        if not os.path.exists(folder):
            return

        for file in os.listdir(folder):

            path = os.path.join(folder, file)

            try:

                if os.path.isfile(path):
                    os.remove(path)

            except Exception:
                pass
            
    # ============================================================
    # Convert PDF to Word
    # ============================================================

    def convert_pdf_to_word(self, pdf_path, output_docx):

        try:

            # Reset Progress
            self.update_progress(0)

            # Create Word Document
            self.create_document()

            # Create Output Folder
            image_folder = self.create_output_folder(output_docx)

            # Remove Old Images
            self.clear_image_folder(image_folder)

            # ------------------------------------------------
            # Extract Text
            # ------------------------------------------------

            print("Extracting Text...")

            text_data = self.extract_text(pdf_path)

            self.write_text(text_data)

            # ------------------------------------------------
            # Extract Tables
            # ------------------------------------------------

            print("Extracting Tables...")

            tables = self.extract_tables(pdf_path)

            self.write_tables(tables)

            # ------------------------------------------------
            # Extract Images
            # ------------------------------------------------

            print("Extracting Images...")

            images = self.extract_images(
                pdf_path,
                image_folder
            )

            self.write_images(images)

            # ------------------------------------------------
            # Save Document
            # ------------------------------------------------

            self.document.save(output_docx)

            self.update_progress(100)

            return True

        except Exception as e:

            print("Conversion Error :", e)

            return False


    # ============================================================
    # Convert Selected Pages
    # ============================================================

    def convert_selected_pages(
            self,
            pdf_path,
            output_docx,
            start_page,
            end_page
    ):

        try:

            self.create_document()

            image_folder = self.create_output_folder(output_docx)

            with pdfplumber.open(pdf_path) as pdf:

                total_pages = len(pdf.pages)

                start_page = max(1, start_page)
                end_page = min(end_page, total_pages)

                for page_no in range(start_page - 1, end_page):

                    page = pdf.pages[page_no]

                    text = page.extract_text()

                    self.document.add_heading(
                        f"Page {page_no + 1}",
                        level=2
                    )

                    if text:

                        self.document.add_paragraph(text)

                    progress = int(
                        ((page_no - start_page + 2) /
                        (end_page - start_page + 1)) * 100
                    )

                    self.update_progress(progress)

            pdf = fitz.open(pdf_path)

            image_index = 1

            for page_no in range(start_page - 1, end_page):

                page = pdf.load_page(page_no)

                for img in page.get_images(full=True):

                    xref = img[0]

                    base = pdf.extract_image(xref)

                    image_bytes = base["image"]

                    ext = base["ext"]

                    filename = os.path.join(
                        image_folder,
                        f"page_{page_no+1}_{image_index}.{ext}"
                    )

                    with open(filename, "wb") as fp:

                        fp.write(image_bytes)

                    self.document.add_picture(
                        filename,
                        width=Inches(5.5)
                    )

                    image_index += 1

            pdf.close()

            self.document.save(output_docx)

            self.update_progress(100)

            return True

        except Exception as e:

            print("Page Conversion Error :", e)

            return False


    # ============================================================
    # PDF Information
    # ============================================================

    def get_pdf_information(self, pdf_path):

        info = {}

        try:

            pdf = fitz.open(pdf_path)

            info["file_name"] = os.path.basename(pdf_path)
            info["file_size"] = os.path.getsize(pdf_path)
            info["total_pages"] = pdf.page_count

            info["title"] = pdf.metadata.get("title", "")
            info["author"] = pdf.metadata.get("author", "")
            info["subject"] = pdf.metadata.get("subject", "")
            info["creator"] = pdf.metadata.get("creator", "")
            info["producer"] = pdf.metadata.get("producer", "")

            info["images"] = self.get_image_count(pdf_path)

            pdf.close()

        except Exception as e:
            print(e)

        return info
        
    # ============================================================
    # Save Word Document
    # ============================================================

    def save_document(self, output_file):

        try:

            if self.document is not None:
                self.document.save(output_file)
                return True

        except Exception as e:

            print("Save Error :", e)

        return False


    # ============================================================
    # Reset Converter
    # ============================================================

    def reset(self):

        self.document = None
        self.progress_callback = None


    # ============================================================
    # Validate PDF
    # ============================================================

    def validate_pdf(self, pdf_path):

        try:

            if not os.path.exists(pdf_path):
                return False

            with pdfplumber.open(pdf_path):
                pass

            return True

        except Exception:

            return False


    # ============================================================
    # Get Total Pages
    # ============================================================

    def get_total_pages(self, pdf_path):

        try:

            with pdfplumber.open(pdf_path) as pdf:
                return len(pdf.pages)

        except Exception:

            return 0


    # ============================================================
    # Extract Text from Single Page
    # ============================================================

    def extract_single_page(self, pdf_path, page_number):

        try:

            with pdfplumber.open(pdf_path) as pdf:

                if page_number < 1 or page_number > len(pdf.pages):
                    return ""

                page = pdf.pages[page_number - 1]

                text = page.extract_text()

                return text if text else ""

        except Exception as e:

            print("Page Extraction Error :", e)

            return ""


    # ============================================================
    # Extract Images Only
    # ============================================================

    def export_images_only(self, pdf_path, output_folder):

        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

        return self.extract_images(
            pdf_path,
            output_folder
        )


    # ============================================================
    # Extract Tables Only
    # ============================================================

    def export_tables_only(self, pdf_path):

        return self.extract_tables(pdf_path)


    # ============================================================
    # Extract Text Only
    # ============================================================

    def export_text_only(self, pdf_path):

        return self.extract_text(pdf_path)


    # ============================================================
    # PDF Summary
    # ============================================================

    def pdf_summary(self, pdf_path):

        info = self.get_pdf_information(pdf_path)

        summary = {

            "Pages": info.get("Pages", 0),
            "Images": info.get("Images", 0),
            "Title": info.get("Title", ""),
            "Author": info.get("Author", ""),
            "Creator": info.get("Creator", ""),
            "Producer": info.get("Producer", "")

        }

        return summary

    # ============================================================
    # Exact Layout PDF To Word Conversion
    # Using pdf2docx
    # ============================================================

    def convert_exact_layout(self, pdf_path, output_docx):

        try:

            self.update_progress(0)

            print("Starting Exact Layout Conversion...")

            cv = Converter(pdf_path)

            cv.convert(
                output_docx,
                start=0,
                end=None
            )

            cv.close()


            self.update_progress(100)

            print("Exact Layout Conversion Completed")

            return True


        except Exception as e:

            print("Exact Layout Conversion Error :", e)

            return False

    # ============================================================
    # Close Converter
    # ============================================================

    def close(self):

        self.reset()


    # ============================================================
    # Destructor
    # ============================================================

    def __del__(self):

        self.close()